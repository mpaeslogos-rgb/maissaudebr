# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

GREEN = RGBColor(0x1B, 0x5E, 0x3F)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading('', level=1)
    run = p.add_run(text)
    run.font.color.rgb = GREEN
    return p

def h2(text):
    p = doc.add_heading('', level=2)
    run = p.add_run(text)
    run.font.color.rgb = GREEN
    return p

def body(text):
    return doc.add_paragraph(text)

def screenshot(label='Inserir print da tela correspondente'):
    p = doc.add_paragraph()
    run = p.add_run(f'[ {label} ]')
    run.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    p.paragraph_format.left_indent = Cm(1)
    return p

def tip(text, prefix='DICA'):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{prefix}: ')
    r1.bold = True
    r1.font.color.rgb = GREEN
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(1)
    return p

# ══════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('MaisSaude BR')
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = GREEN

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('Sistema de Gestao de Clinica Medica').font.size = Pt(16)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = m.add_run('Manual do Operador  |  Versao 1.0  |  Maio 2026')
mr.italic = True; mr.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 1. DESCRICAO
# ══════════════════════════════════════════════════════
h1('1. Descricao do Sistema')
body('O MaisSaude BR e uma plataforma web completa para gestao de clinicas medicas. '
     'Centraliza agenda, pacientes, medicos, prontuarios, financeiro e WhatsApp em '
     'um unico sistema integrado, acessivel de qualquer dispositivo pelo navegador.')
h2('Modulos do sistema')
for m in ['Agenda', 'Pacientes', 'Medicos', 'Prontuarios',
          'Financeiro (A Receber / A Pagar / Fluxo de Caixa)',
          'WhatsApp / Chats', 'Relatorios', 'Configuracoes', 'Auditoria']:
    doc.add_paragraph(m, style='List Bullet')
h2('Tecnologia')
body('Sistema 100% web. Recomendado: Google Chrome ou Microsoft Edge, versao mais recente. '
     'Nao requer instalacao. Funciona em computadores, tablets e smartphones.')
screenshot('Print da tela principal do sistema')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 2. LOGIN
# ══════════════════════════════════════════════════════
h1('2. Acesso e Login')
h2('Como fazer login')
for step in [
    'Abra o navegador e acesse o endereco do sistema.',
    'Informe seu e-mail cadastrado.',
    'Informe sua senha. Clique no icone de olho para ver o que esta digitando.',
    'Clique em "Entrar". O sistema exibe o painel inicial.',
]:
    doc.add_paragraph(step, style='List Number')
screenshot('Print da tela de login')
h2('Esqueci minha senha')
body('Clique em "Esqueci minha senha". Uma mensagem orientara a entrar em contato com o administrador.')
h2('Encerrar sessao (Logout)')
body('Clique no menu do usuario (canto superior direito) e selecione "Sair". '
     'A sessao expira automaticamente apos 8 horas.')
tip('Sempre faca logout ao sair, especialmente em equipamentos compartilhados.', 'ATENCAO')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 3. PERFIS
# ══════════════════════════════════════════════════════
h1('3. Perfis de Usuario (Roles)')
body('O sistema possui 4 perfis de acesso. Cada usuario e cadastrado com um perfil '
     'que define o que pode visualizar e fazer.')
for perfil, desc in [
    ('ADMIN', 'Acesso total: todos os modulos, auditoria e cadastro de usuarios.'),
    ('MEDICO', 'Agenda, prontuarios, pacientes e relatorios.'),
    ('RECEPCIONISTA', 'Agenda, pacientes, financeiro basico e WhatsApp.'),
    ('PACIENTE', 'Portal do paciente (acesso futuro).'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{perfil}: ').bold = True
    p.add_run(desc)
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 4. DASHBOARD
# ══════════════════════════════════════════════════════
h1('4. Dashboard - Painel Inicial')
body('O Dashboard e a primeira tela apos o login. Exibe um resumo operacional e financeiro.')
h2('Indicadores (KPIs)')
for kpi in [
    'Total de Pacientes: numero total cadastrado no sistema.',
    'Consultas Hoje: quantidade de consultas agendadas para o dia atual.',
    'Faturamento Mensal: soma de pagamentos confirmados no mes corrente.',
    'Contas a Pagar: total de despesas pendentes.',
]:
    doc.add_paragraph(kpi, style='List Bullet')
h2('Proximas Consultas')
body('Lista das 5 proximas consultas agendadas/confirmadas. Clique em qualquer uma para ver detalhes.')
screenshot('Print do Dashboard')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 5. AGENDA
# ══════════════════════════════════════════════════════
h1('5. Agenda - Consultas')
body('A Agenda e o coracao do sistema. Permite visualizar, criar, confirmar e concluir consultas.')

h2('Navegar pela Agenda')
body('Use os botoes < > para navegar entre semanas/meses. O botao "Hoje" retorna ao periodo atual. '
     'Alterne entre Visao Semanal (grade horaria 8h-20h) e Visao Mensal (calendario). '
     'Filtre por medico no seletor superior.')
screenshot('Print da Agenda Semanal')

h2('Agendar Nova Consulta')
for step in [
    'Clique em "+ Nova Consulta" ou clique em um horario vazio na grade semanal.',
    'Selecione o Paciente no menu suspenso (mostra nome e CPF).',
    'Selecione o Medico. O campo "Valor da consulta" e preenchido automaticamente com o valor cadastrado do medico.',
    'Informe a Data e o Horario de inicio (hora e minuto).',
    'Escolha a Duracao (15 min a 2 horas).',
    'Opcional: edite o Valor se houver negociacao de preco diferente do padrao.',
    'Opcional: adicione um Motivo da consulta.',
    'Clique em "Agendar Consulta".',
]:
    doc.add_paragraph(step, style='List Number')
tip('O valor digitado vai automaticamente para o Contas a Receber do modulo Financeiro.')
screenshot('Print do formulario Nova Consulta')

h2('Status das Consultas')
for status, desc in [
    ('Agendada', 'Registrada, aguardando confirmacao. Acoes: Confirmar, Cancelar, No-show.'),
    ('Confirmada', 'Paciente confirmou. Acoes: Iniciar, Cancelar, No-show.'),
    ('Em Andamento', 'Consulta em curso. Acoes: Finalizar, Cancelar.'),
    ('Concluida', 'Encerrada com sucesso. Estado terminal.'),
    ('Cancelada', 'Cancelada. Estado terminal.'),
    ('No-show', 'Paciente nao compareceu. Estado terminal.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{status}: ').bold = True
    p.add_run(desc)

h2('Teleconsulta (Videochamada)')
for step in [
    'Com consulta Confirmada, clique em "Iniciar Consulta" para gerar a sala de video.',
    'Copie o link e envie ao paciente pelo WhatsApp ou e-mail.',
    'O video abre em janela flutuante (Picture-in-Picture) para o medico acompanhar.',
    'Ao concluir, clique em "Finalizar".',
]:
    doc.add_paragraph(step, style='List Number')
screenshot('Print do Painel de Detalhes / Prontuario')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 6. PACIENTES
# ══════════════════════════════════════════════════════
h1('6. Pacientes')
body('Gestao completa da base de pacientes da clinica.')
h2('Buscar Paciente')
for step in [
    'Na barra de busca, digite nome, CPF ou e-mail.',
    'A lista filtra automaticamente em tempo real.',
    'Para limpar a busca, clique no X ao lado do campo.',
]:
    doc.add_paragraph(step, style='List Number')

h2('Cadastrar Novo Paciente')
for step in [
    'Clique em "+ Novo Paciente".',
    'Preencha: Nome completo, CPF (formatado automaticamente), Data de nascimento, Genero, Telefone.',
    'Opcionais: E-mail, Endereco, Observacoes (alergias, condicoes, etc.).',
    'Clique em "Salvar".',
]:
    doc.add_paragraph(step, style='List Number')

h2('Excluir / Anonimizar Paciente (LGPD)')
body('Por exigencia da LGPD (Lei 13.709/2018), ao excluir um paciente o sistema '
     'realiza a anonimizacao dos dados. O registro permanece para rastreabilidade, '
     'mas todos os dados pessoais sao removidos. Esta acao nao pode ser desfeita.')
tip('Importe pacientes em lote clicando em "Importar Excel" e selecionando um arquivo .xlsx.')
screenshot('Print da lista de Pacientes e formulario de cadastro')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 7. MEDICOS
# ══════════════════════════════════════════════════════
h1('7. Medicos')
body('Cadastro e gestao dos profissionais de saude da clinica.')
h2('Cadastrar Novo Medico')
for step in [
    'Clique em "+ Novo Medico".',
    'Preencha: Nome completo, E-mail, Senha de acesso, CRM, Estado do CRM, Especialidade.',
    'Opcionais: Telefone, Valor padrao da consulta (R$), Bio/Apresentacao.',
    'Clique em "Cadastrar". Um usuario de acesso e criado automaticamente.',
]:
    doc.add_paragraph(step, style='List Number')
tip('O Valor padrao da consulta e preenchido automaticamente no agendamento e pode ser editado por negociacao.')
h2('Editar e Desativar Medico')
body('Clique no icone de lapis para editar especialidade, telefone, valor e bio. '
     'Clique em "Desativar" para impedir novos agendamentos (historico e preservado).')
screenshot('Print da lista de Medicos e formulario de cadastro')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 8. PRONTUARIOS
# ══════════════════════════════════════════════════════
h1('8. Prontuarios Medicos')
body('Registro eletronico de consultas. Exclusivo para medicos e administradores.')
h2('Preencher Prontuario durante a Consulta')
for step in [
    'Na Agenda, clique na consulta e abra a aba "Prontuario".',
    'Preencha a Queixa Principal.',
    'Registre a Historia da Doenca. Use os modelos por especialidade para agilizar.',
    'Informe o Diagnostico usando a busca integrada de CID-10.',
    'Registre a Conduta / Prescricao e Observacoes.',
    'Expanda Sinais Vitais (pressao, FC, temperatura, SpO2, peso, altura) se necessario.',
    'Clique em "Salvar".',
]:
    doc.add_paragraph(step, style='List Number')
h2('Imprimir Prescricao')
body('Clique em "Imprimir Receita" para gerar um PDF formatado com dados do paciente, medico, CRM e prescricao.')
tip('Prontuarios NAO podem ser excluidos. Retencao minima de 20 anos (CFM 1.821/2007).', 'ATENCAO')
screenshot('Print da aba Prontuario')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 9. FINANCEIRO
# ══════════════════════════════════════════════════════
h1('9. Financeiro')
body('Controle completo de receitas, despesas e fluxo de caixa. Tres abas integradas.')
h2('Aba 1 - Fluxo de Caixa')
body('Visao consolidada: saldo realizado, entradas e saidas do mes, alertas de vencimento, '
     'grafico de evolucao mensal (3/6/12/24 meses) e projecao dos proximos 30 dias.')
screenshot('Print do Fluxo de Caixa')

h2('Aba 2 - Contas a Receber')
body('Cobranças geradas automaticamente ao agendar consultas. '
     'Filtros: Todos, Pendente, Vencido, Recebido, Cancelado.')
for step in [
    'Localize a cobranca na lista.',
    'Clique em "Receber" para marcar como recebido.',
    'O status muda para Recebido e o valor e contabilizado no Fluxo de Caixa.',
]:
    doc.add_paragraph(step, style='List Number')
screenshot('Print do Contas a Receber')

h2('Aba 3 - Contas a Pagar')
body('Controle de despesas: aluguel, energia, fornecedores, etc.')
for step in [
    'Clique em "+ Nova Despesa".',
    'Informe: Descricao (obrigatorio), Categoria, Fornecedor, Valor (R$) e Vencimento.',
    'Clique em "Salvar".',
    'Para pagar: clique em "Pagar" na linha. Para cancelar: clique no icone de lixeira.',
]:
    doc.add_paragraph(step, style='List Number')
tip('Clique em "Atualizar" para recarregar os dados. A pagina nao atualiza automaticamente.')
screenshot('Print do Contas a Pagar')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 10. WHATSAPP
# ══════════════════════════════════════════════════════
h1('10. Chats / WhatsApp')
body('Atendimento integrado pelo WhatsApp com IA automatica e handoff manual para a equipe.')
h2('Responder Manualmente')
body('Clique na conversa na barra lateral. Digite no campo inferior e pressione Enter ou clique em Enviar.')
h2('Pausar / Retomar a IA')
for step in [
    'Com a conversa aberta, clique em "Pausar IA".',
    'A IA para de responder naquele chat.',
    'Para reativar, clique em "Retomar IA".',
]:
    doc.add_paragraph(step, style='List Number')
h2('Assumir Atendimento (Duplo Clique)')
body('Duplo clique em uma conversa para abrir em janela dedicada com IA pausada. '
     'Responda normalmente. Clique em "Devolver" para retornar a conversa a IA.')
h2('Mensagem em Massa')
for step in [
    'Clique em "Mensagem em Massa".',
    'Digite o texto (limite: 4.096 caracteres).',
    'Clique em "Continuar" e revise a confirmacao.',
    'Clique em "Enviar para Todos". A mensagem sera enviada a todos os pacientes.',
]:
    doc.add_paragraph(step, style='List Number')
tip('O envio em massa e irreversivel. Revise o conteudo com cuidado antes de confirmar.', 'ATENCAO')
screenshot('Print do modulo de Chats')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 11. RELATORIOS
# ══════════════════════════════════════════════════════
h1('11. Relatorios')
body('KPIs e graficos para analise de desempenho da clinica.')
for kpi in [
    'Total de Pacientes',
    'Total de Consultas',
    'Faturamento (pagamentos confirmados)',
    'Resultado Liquido (faturamento - despesas)',
    'Taxa de Conversao (% consultas concluidas)',
    'Inadimplencia (cobranças em aberto)',
]:
    doc.add_paragraph(kpi, style='List Bullet')
h2('Exportar Relatorio')
body('Clique em "Exportar Relatorio". Use Ctrl+P no navegador e selecione "Salvar como PDF".')
screenshot('Print da tela de Relatorios')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 12. CONFIGURACOES
# ══════════════════════════════════════════════════════
h1('12. Configuracoes')
body('Parametrizacao da clinica e das mensagens automaticas do WhatsApp. Acesso: ADMIN.')
for sec, desc in [
    ('Dados da Clinica', 'Nome, numero WhatsApp, horarios, especialidades.'),
    ('Telemedicina', 'Plataforma, instrucoes de acesso, link de consulta.'),
    ('Pagamentos', 'Valor padrao, chave PIX, link de pagamento, metodos aceitos.'),
    ('Politicas da Clinica', 'Confirmacao, cancelamento e remarcacao.'),
    ('Mensagens Automaticas', 'Boas-vindas, confirmacao, lembrete, cobrança, emergencia.'),
    ('Regras da IA', 'Instrucoes adicionais para o comportamento da IA.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{sec}: ').bold = True
    p.add_run(desc)
body('Para salvar: edite os campos necessarios e clique em "Salvar Configuracoes" no rodape.')
screenshot('Print da tela de Configuracoes')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# 13. AUDITORIA
# ══════════════════════════════════════════════════════
h1('13. Auditoria (Admin)')
body('Rastreabilidade completa de todas as acoes do sistema. Acesso: ADMIN.')
for acao, desc in [
    ('LOGIN / LOGIN_FAILED', 'Entradas e tentativas falhas de acesso.'),
    ('CREATE', 'Novos cadastros (paciente, medico, consulta, etc.).'),
    ('UPDATE', 'Edicoes em registros existentes.'),
    ('DELETE / CANCEL', 'Exclusoes e cancelamentos.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{acao}: ').bold = True
    p.add_run(desc)
h2('Filtrar e Ver Detalhes')
for step in [
    'Use a busca por texto para encontrar registros especificos.',
    'Filtre por Entidade (Paciente, Consulta, etc.) e por Acao.',
    'Defina um intervalo de datas.',
    'Clique no botao de expansao (>) para ver os metadados em JSON.',
]:
    doc.add_paragraph(step, style='List Number')
screenshot('Print da tela de Auditoria')
doc.add_page_break()

# ══════════════════════════════════════════════════════
# APENDICE A
# ══════════════════════════════════════════════════════
h1('Apendice A - Fluxo Completo de uma Consulta')
for i, step in enumerate([
    'Paciente entra em contato pelo WhatsApp (IA / Recepcao)',
    'Recepcao cadastra o paciente se for novo (modulo Pacientes)',
    'Recepcao agenda a consulta com medico e valor (modulo Agenda)',
    'Sistema gera cobranca automaticamente no Contas a Receber',
    'Recepcao confirma a presenca do paciente na Agenda',
    'Medico inicia a consulta e preenche o prontuario',
    'Medico finaliza a consulta',
    'Recepcao registra o pagamento no Contas a Receber',
    'Sistema atualiza o Fluxo de Caixa automaticamente',
], 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# APENDICE B - FAQ
# ══════════════════════════════════════════════════════
h1('Apendice B - Perguntas Frequentes')
faqs = [
    ('O sistema funciona em celular?',
     'Sim, e responsivo. Recomendamos tablets ou computadores para melhor experiencia na Agenda e Financeiro.'),
    ('Posso acessar de casa?',
     'Sim, o sistema e 100% web. Acesse pelo navegador de qualquer dispositivo com internet.'),
    ('Posso agendar uma consulta no passado?',
     'Nao. O horario deve ser no futuro. Para registrar consultas passadas, use Prontuarios diretamente.'),
    ('Por que a cobranca nao aparece no Contas a Receber?',
     'A cobranca e gerada apenas quando o medico tem valor cadastrado OU quando o valor e digitado no formulario de agendamento.'),
    ('Como cancelar uma cobranca incorreta?',
     'No momento, entre em contato com o administrador para correcao direta. A funcao de estorno sera adicionada em versao futura.'),
    ('A IA nao esta respondendo. O que fazer?',
     'Verifique em Configuracoes se o numero e credenciais Z-API estao corretos. Consulte o administrador.'),
    ('Posso editar uma mensagem enviada pelo WhatsApp?',
     'Nao. Mensagens enviadas nao podem ser editadas. Envie uma nova mensagem com a correcao.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
# RODAPE
# ══════════════════════════════════════════════════════
doc.add_page_break()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run('MaisSaude BR  |  Manual do Operador v1.0  |  Maio 2026')
fr.italic = True
fr.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save(r'C:\Users\Marcos Paes\OneDrive\Desktop\maissaudebr\docs\Manual_MaisSaude_BR.docx')
print('DOCX gerado com sucesso!')
